from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.first_line_indent = Cm(0.74)
pf.space_after = Pt(4)

desktop = r'D:\Documents\Desktop'

# === 标题 ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.first_line_indent = Cm(0)
run = title.add_run('RTX 5090 装机配置清单')
run.bold = True
run.font.size = Pt(16)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.first_line_indent = Cm(0)
run = sub.add_run('石家庄·太和电子城  |  仅配主机  |  2026年6月')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# === 一、推荐配置 ===
def add_h(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_para(doc, text, bold=False, color=None, indent=True):
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

add_h(doc, '一、推荐配置单（AI 本地模型专用）')

table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'

headers = ['配件', '型号', '参考价格']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

data = [
    ['🔥 显卡', 'RTX 5090 D 32GB GDDR7', '¥21,999'],
    ['🧠 CPU', 'Intel Core Ultra 9 285K', '¥4,799'],
    ['🔌 主板', '华硕 Z890-PLUS WIFI', '¥2,799'],
    ['💾 内存', 'DDR5 6800 64GB(32Gx2)', '¥779'],
    ['📀 固态', '三星 990 PRO 2TB', '¥1,119'],
    ['🔋 电源', '1200W 白金 ATX 3.0', '¥1,099'],
    ['❄️ 散热', '360mm 一体水冷', '¥799'],
    ['📦 机箱', '大塔机箱(支持380mm显卡)', '¥599'],
]

for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(val)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if j == 2:
            run.bold = True

row = table.add_row()
cells = row.cells
cells[0].merge(cells[1])
p = cells[0].paragraphs[0]
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('💰 合计')
run.bold = True
run.font.size = Pt(15)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = cells[2].paragraphs[0]
p2.paragraph_format.first_line_indent = Cm(0)
run2 = p2.add_run('≈ ¥34,000')
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run2.font.name = '宋体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# === 二、显卡价格 ===
add_h(doc, '二、显卡价格参考')

for g in ['RTX 5090 D（中国特供版）          ¥21,999',
          'RTX 5090 公版 / 非公版                 ¥24,399',
          'ROG RTX 5090 白夜神 OC                 ¥26,300']:
    add_para(doc, g)

doc.add_paragraph()

# === 三、太和电子城 ===
add_h(doc, '三、石家庄·太和电子城购买建议')

tips = [
    ('📍 地址：', '新华区民族路 太和电子城，推荐 3-4 楼 DIY 装机区'),
    ('💡 比价：', '多问几家，整机砍价空间约 5-10%（约省 2000-3000 元）'),
    ('🔍 注意：', '问清楚是 5090 D 还是海外版，差价好几千'),
]
for label, content in tips:
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = '宋体'
    r1.font.size = Pt(14)
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r2 = p.add_run(content)
    r2.font.name = '宋体'
    r2.font.size = Pt(14)
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# === 四、注意事项 ===
add_h(doc, '四、装机注意事项（重点）')

for w in ['电源至少 1200W 白金，必须支持 ATX 3.0 + 原生 12VHPWR 接口',
          '机箱要大，5090 非公版长度 330-385mm，务必确认放得下',
          '买 90 度弯头电源线，避免 12VHPWR 线弯折烧毁',
          '散热要 360mm 水冷，5090 发热量大，机箱风道要好',
          '5090 D 为中国特供版，性能几乎无差别但价格低几千，性价比最高']:
    p = doc.add_paragraph()
    run = p.add_run('▸ ' + w)
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# === 五、本地模型 ===
add_h(doc, '五、装好后能跑的本地开源模型')

for name, desc in [
    ('Qwen3.6-27B（阿里通义）-- 约 40 tok/s', '中文综合最强，262K 上下文，推荐首选'),
    ('DeepSeek R1 32B（深度求索）-- 约 35 tok/s', '编程推理最强，代码党必备'),
    ('Qwen3.6 35B-A3B MoE -- 超 100 tok/s', '极速推理，仅激活 3B 参数'),
    ('Gemma 4 31B（Google）-- 约 45 tok/s', '131K 长上下文，多模态图文音频'),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.name = '宋体'
    r1.font.size = Pt(14)
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p2 = doc.add_paragraph()
    r2 = p2.add_run(desc)
    r2.font.name = '宋体'
    r2.font.size = Pt(14)
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run('推荐工具：LM Studio（图形界面）或 Ollama（命令行）')
r.bold = True
r.font.name = '宋体'
r.font.size = Pt(14)
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页脚
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.first_line_indent = Cm(0)
run = footer.add_run('-- 生成于 2026年6月20日 价格仅供参考 --')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

filepath = os.path.join(desktop, 'RTX5090-装机配置清单.docx')
doc.save(filepath)
print(f'OK: {filepath}')
